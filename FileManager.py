import os
import docx
import random
from YesNo import YesNo
from docx2pdf import convert
from Values import MainValues
import NotBaseValues
values = MainValues()

def CreateFolders():
    listOfFiles = os.listdir()
    for i in range(1, 8):
        if not os.path.isdir(f"ROUND{i}"):
            os.mkdir(f"ROUND{i}")
        else:
            listOfFiles = os.listdir(f"ROUND{i}")
            for file in listOfFiles:
                try:
                    os.remove(f"ROUND{i}"+"/"+file)
                except: None

def Upload(listOfCountries: list, ecology: float, values: dict, round: int):
    for country in listOfCountries:
        doc = docx.Document("Examples/"+country.name+'Exp.docx')
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        style = doc.styles['Normal']
                        font = style.font
                        font.size = docx.shared.Pt(14)
                        for i in range(len(country.listOfTowns)):
                            if country.listOfTowns[i].existance:
                                paragraph.text = paragraph.text.replace(f"Develop{i+1}", str(country.listOfTowns[i].development))
                                paragraph.text = paragraph.text.replace(f"Qual{i+1}", str(country.listOfTowns[i].lifeQual))
                                paragraph.text = paragraph.text.replace(f"Profit{i+1}", str(country.listOfTowns[i].profit))
                                paragraph.text = paragraph.text.replace(f"Shield{i+1}", YesNo(country.listOfTowns[i].shield))
                                paragraph.text = paragraph.text.replace(f"PVO{i + 1}",
                                                                        YesNo(country.listOfTowns[i].PVO))
                            else:
                                paragraph.text = paragraph.text.replace(f"Develop{i + 1}%",
                                                                        " ")
                                paragraph.text = paragraph.text.replace(f"Qual{i + 1}%",
                                                                        " ")
                                paragraph.text = paragraph.text.replace(f"Profit{i + 1}$",
                                                                        " ")
                                paragraph.text = paragraph.text.replace(f"Shield{i + 1}",
                                                                        " ")
                                paragraph.text = paragraph.text.replace(f"PVO{i + 1}",
                                                                        " ")
                        paragraph.text = paragraph.text.replace('avg', str(country.avgLifeQual))
                        paragraph.text = paragraph.text.replace('budget', str(country.budget))
                        paragraph.text = paragraph.text.replace('nuke', YesNo(country.tech))
                        paragraph.text = paragraph.text.replace('nke', str(country.bombAmountNow))
                        paragraph.text = paragraph.text.replace('scouts', str(country.scoutsHome))
                        paragraph.text = paragraph.text.replace('inform', country.message)
        filePath = country.name+str(round)+'.docx'
        try:
            doc.save("ROUND"+str(round)+"\\"+ filePath)
        except:
            filePath = country.name+str(round)+"_"+str(random.randint(1,100000))+'.docx'
            doc.save("ROUND"+str(round)+"\\"+ filePath)
        convert("ROUND"+str(round)+"\\"+ filePath)
    doc = docx.Document("Examples/RoundResultExp.docx")
    resEc = int(ecology / values["begEcology"] * 100)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = paragraph.text.replace('Экология  [EcAmount%]', f"Экология  [{resEc}%]")
                    style = doc.styles['Normal']
                    font = style.font
                    font.size = docx.shared.Pt(13)
                    for i in range(len(listOfCountries)):
                        paragraph.text = paragraph.text.replace(f"{i}CountryName", str(listOfCountries[i].name))
                        paragraph.text = paragraph.text.replace(f"{i}CountryPer",
                                                                str(listOfCountries[i].avgLifeQual))
                        for z in range(len(listOfCountries[i].listOfTowns)):
                            paragraph.text = paragraph.text.replace(f"{i}{z}CityName", str(listOfCountries[i].listOfTowns[z].name))
                            paragraph.text = paragraph.text.replace(f"{i}{z}CityPer", str(listOfCountries[i].listOfTowns[z].lifeQual))
    filePath = "RoundResult"+str(round)+'.docx'
    doc.save("ROUND"+str(round)+"\\"+ filePath)
    convert("ROUND"+str(round)+"\\"+ filePath)